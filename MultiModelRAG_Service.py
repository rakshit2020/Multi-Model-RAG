
from PIL import Image
import io
import matplotlib.pyplot as plt
import os
import uuid
import base64
import re
from typing import List, Dict, Any, Optional
import pandas as pd
import docx
from pptx import Presentation
import csv
import tempfile
from unstructured.partition.pdf import partition_pdf
from unstructured.partition.docx import partition_docx
from unstructured.partition.pptx import partition_pptx
from langchain_groq import ChatGroq
from langchain_cohere import CohereEmbeddings
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough, RunnableLambda
from langchain_core.messages import HumanMessage
from langchain_chroma import Chroma
from langchain.storage import InMemoryStore
from langchain.schema.document import Document
from langchain.retrievers.multi_vector import MultiVectorRetriever
from langchain.text_splitter import RecursiveCharacterTextSplitter


class MultimodalRAGProcessor:
    def __init__(self, groq_api_key: str = None, cohere_api_key: str = None):
        if groq_api_key:
            os.environ["GROQ_API_KEY"] = groq_api_key
        if cohere_api_key:
            os.environ["COHERE_API_KEY"] = cohere_api_key

        self.llm = ChatGroq(
            temperature=0.3,
            model="meta-llama/llama-4-scout-17b-16e-instruct"
        )

        self.embeddings = CohereEmbeddings(
            model="embed-multilingual-v3.0",
            cohere_api_key=os.environ.get("COHERE_API_KEY")
        )

        self.vectorstore = Chroma(
            collection_name="enhanced_multi_modal_rag",
            embedding_function=self.embeddings
        )
        self.store = InMemoryStore()
        self.id_key = "doc_id"

        self.retriever = MultiVectorRetriever(
            vectorstore=self.vectorstore,
            docstore=self.store,
            id_key=self.id_key,
        )

        self.supported_extensions = {'.pdf', '.docx', '.pptx', '.xlsx', '.txt', '.csv'}

        self._setup_chains()

    def _setup_chains(self):
        prompt_text = """
        You are an assistant tasked with summarizing content from various document types.
        Give a concise but comprehensive summary that captures the key information.
        Important guidelines:
        - Focus on the main concepts, findings, or data presented
        - Preserve technical terms and important details
        - Keep the summary informative but concise
        - Do not add your own interpretations

        Content to summarize: {element}

        Summary:"""

        self.summary_prompt = ChatPromptTemplate.from_template(prompt_text)
        self.summarize_chain = (
                {"element": lambda x: x}
                | self.summary_prompt
                | self.llm
                | StrOutputParser()
        )

    def extract_document_elements(self, file_path: str, file_type: str = None):
        if file_type is None:
            file_type = os.path.splitext(file_path)[1].lower()

        print(f"Extracting elements from {file_type} file")

        if file_type == '.pdf':
            return self._extract_pdf_elements(file_path)
        elif file_type == '.docx':
            return self._extract_docx_elements(file_path)
        elif file_type == '.pptx':
            return self._extract_pptx_elements(file_path)
        elif file_type == '.xlsx':
            return self._extract_xlsx_elements(file_path)
        elif file_type == '.txt':
            return self._extract_txt_elements(file_path)
        elif file_type == '.csv':
            return self._extract_csv_elements(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def _extract_pdf_elements(self, file_path: str):
        chunks = partition_pdf(
            filename=file_path,
            infer_table_structure=True,
            strategy="hi_res",
            extract_image_block_types=["Image"],
            extract_image_block_to_payload=True,
            chunking_strategy="by_title",
            max_characters=10000,
            combine_text_under_n_chars=2000,
            new_after_n_chars=6000,
        )
        print(f"Extracted {len(chunks)} chunks from PDF")
        return chunks

    def _extract_docx_elements(self, file_path: str):
        chunks = partition_docx(
            filename=file_path,
            chunking_strategy="by_title",
            max_characters=10000,
            combine_text_under_n_chars=2000,
            new_after_n_chars=6000,
        )
        print(f"Extracted {len(chunks)} chunks from DOCX")
        return chunks

    def _extract_pptx_elements(self, file_path: str):
        chunks = partition_pptx(
            filename=file_path,
            chunking_strategy="by_title",
            max_characters=10000,
            combine_text_under_n_chars=2000,
            new_after_n_chars=6000,
        )
        print(f"Extracted {len(chunks)} chunks from PPTX")
        return chunks

    def _extract_xlsx_elements(self, file_path: str):
        chunks = []
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                # Convert dataframe to text representation
                text_content = f"Sheet: {sheet_name}\n"
                text_content += df.to_string(index=False)

                # Create a document-like object
                chunk = type('Chunk', (), {
                    'text': text_content,
                    'metadata': type('Metadata', (), {
                        'sheet_name': sheet_name,
                        'file_type': 'xlsx'
                    })()
                })()
                chunks.append(chunk)

        except Exception as e:
            print(f"Error reading Excel file: {e}")

        print(f"Extracted {len(chunks)} sheets from XLSX")
        return chunks

    def _extract_txt_elements(self, file_path: str):
        chunks = []
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()

            # Calculate average paragraph size for dynamic chunking
            avg_paragraph_size = self._calculate_average_paragraph_size(content)

            chunk_size = max(avg_paragraph_size * 3, 500)  # Minimum 500 chars
            chunk_overlap = avg_paragraph_size

            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                length_function=len,
                separators=[
                    "\n\n",  # Double newlines (paragraphs)
                    "\n",  # Single newlines
                    ". ",  # Sentence endings
                    " ",  # Spaces
                    ""  # Characters
                ]
            )

            text_chunks = text_splitter.split_text(content)

            for i, text_chunk in enumerate(text_chunks):
                chunk = type('Chunk', (), {
                    'text': text_chunk.strip(),
                    'metadata': type('Metadata', (), {
                        'chunk_id': i,
                        'file_type': 'txt',
                    })()
                })()
                chunks.append(chunk)

        except Exception as e:
            print(f"Error reading TXT file: {e}")

        return chunks

    def _calculate_average_paragraph_size(self, content: str) -> int:

        try:
            # Split content by double newlines to identify paragraphs
            paragraphs = re.split(r'\n\s*\n', content.strip())

            # Filter out empty paragraphs
            paragraphs = [p.strip() for p in paragraphs if p.strip()]

            if not paragraphs:
                return 1000  # Default fallback

            # Calculate average paragraph length
            total_chars = sum(len(paragraph) for paragraph in paragraphs)
            avg_size = total_chars // len(paragraphs)

            # Ensure reasonable bounds (between 200 and 2000 characters)
            avg_size = max(200, min(avg_size, 2000))

            print(f"Detected {len(paragraphs)} paragraphs, average size: {avg_size} characters")

            return avg_size

        except Exception as e:
            print(f"Error calculating paragraph size: {e}")
            return 1000

    def _extract_csv_elements(self, file_path: str):
        chunks = []
        try:
            df = pd.read_csv(file_path)

            # Convert dataframe to text representation
            text_content = "CSV Data:\n"
            text_content += df.to_string(index=False)

            # Also create summary statistics if numeric columns exist
            numeric_columns = df.select_dtypes(include=['number']).columns
            if len(numeric_columns) > 0:
                text_content += "\n\nSummary Statistics:\n"
                text_content += df[numeric_columns].describe().to_string()

            chunk = type('Chunk', (), {
                'text': text_content,
                'metadata': type('Metadata', (), {
                    'rows': len(df),
                    'columns': len(df.columns),
                    'file_type': 'csv'
                })()
            })()
            chunks.append(chunk)

        except Exception as e:
            print(f"Error reading CSV file: {e}")

        print(f"Extracted {len(chunks)} chunks from CSV")
        return chunks

    def separate_elements(self, chunks):

        print("Separating elements...")

        tables = []
        texts = []

        for chunk in chunks:
            if "Table" in str(type(chunk)):
                tables.append(chunk)
            else:
                texts.append(chunk)

        images = self._get_images_base64(chunks)

        print(f"Found {len(texts)} text chunks, {len(tables)} tables, {len(images)} images")
        return texts, tables, images

    def _get_images_base64(self, chunks):
        images_b64 = []
        for chunk in chunks:
            if hasattr(chunk, 'metadata') and hasattr(chunk.metadata, 'orig_elements'):
                chunk_els = chunk.metadata.orig_elements
                for el in chunk_els:
                    if "Image" in str(type(el)):
                        if hasattr(el.metadata, 'image_base64') and el.metadata.image_base64:
                            images_b64.append(el.metadata.image_base64)
        return images_b64

    def generate_summaries(self, texts, tables, images):
        print("Generating summaries...")

        print("Summarizing text chunks...")
        text_summaries = []
        for text in texts:
            try:
                content = text.text if hasattr(text, 'text') else str(text)
                summary = self.summarize_chain.invoke(content)
                text_summaries.append(summary)
            except Exception as e:
                print(f"Error summarizing text: {e}")
                text_summaries.append("Summary not available")

        print("Summarizing tables...")
        table_summaries = []
        for table in tables:
            try:
                if hasattr(table.metadata, 'text_as_html'):
                    content = table.metadata.text_as_html
                else:
                    content = table.text if hasattr(table, 'text') else str(table)
                summary = self.summarize_chain.invoke(content)
                table_summaries.append(summary)
            except Exception as e:
                print(f"Error summarizing table: {e}")
                table_summaries.append("Table summary not available")

        print("Summarizing images...")
        image_summaries = self._summarize_images_with_groq(images)

        return text_summaries, table_summaries, image_summaries

    def _summarize_images_with_groq(self, images):
        if not images:
            return []

        vision_llm = ChatGroq(
            model="meta-llama/llama-4-scout-17b-16e-instruct",
            temperature=0.3
        )

        image_summaries = []
        for image_b64 in images:
            try:
                prompt_template = """Analyze this image in detail. Describe:
                - What type of visualization, diagram, or content this is
                - Key elements, components, or data shown
                - Any text, labels, or annotations visible
                - The main concept or information being conveyed

                Be specific and detailed in your description."""

                messages = [
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt_template},
                            {
                                "type": "image_url",
                                "image_url": {"url": f"data:image/jpeg;base64,{image_b64}"}
                            }
                        ]
                    }
                ]

                response = vision_llm.invoke(messages)
                image_summaries.append(response.content)

            except Exception as e:
                print(f"Error processing image: {e}")
                image_summaries.append("Image analysis not available")

        return image_summaries

    def load_to_vectorstore(self, texts, tables, images, text_summaries, table_summaries, image_summaries):

        print("Loading data to vectorstore...")

        def clean_summary(summary):
            if not summary or not summary.strip():
                return "No content available"
            return summary.strip()

        if texts and text_summaries:
            doc_ids = [str(uuid.uuid4()) for _ in texts]
            summary_texts = [
                Document(
                    page_content=clean_summary(text_summaries[i]),
                    metadata={self.id_key: doc_ids[i]}
                )
                for i in range(min(len(texts), len(text_summaries)))
                if clean_summary(text_summaries[i])
            ]
            if summary_texts:
                self.retriever.vectorstore.add_documents(summary_texts)
                valid_texts = texts[:len(summary_texts)]
                self.retriever.docstore.mset(list(zip(doc_ids[:len(valid_texts)], valid_texts)))

        if tables and table_summaries:
            table_ids = [str(uuid.uuid4()) for _ in tables]
            summary_tables = [
                Document(
                    page_content=clean_summary(table_summaries[i]),
                    metadata={self.id_key: table_ids[i]}
                )
                for i in range(min(len(tables), len(table_summaries)))
                if clean_summary(table_summaries[i])
            ]
            if summary_tables:
                self.retriever.vectorstore.add_documents(summary_tables)
                valid_tables = tables[:len(summary_tables)]
                self.retriever.docstore.mset(list(zip(table_ids[:len(valid_tables)], valid_tables)))

        if images and image_summaries:
            img_ids = [str(uuid.uuid4()) for _ in images]
            summary_img = [
                Document(
                    page_content=clean_summary(image_summaries[i]),
                    metadata={self.id_key: img_ids[i]}
                )
                for i in range(min(len(images), len(image_summaries)))
                if clean_summary(image_summaries[i])
            ]
            if summary_img:
                self.retriever.vectorstore.add_documents(summary_img)
                valid_images = images[:len(summary_img)]
                self.retriever.docstore.mset(list(zip(img_ids[:len(valid_images)], valid_images)))

        print("Data loaded successfully!")

    def _parse_docs(self, docs):
        b64 = []
        text = []
        for doc in docs:
            if hasattr(doc, 'text'):
                text_content = doc.text
            elif hasattr(doc, 'page_content'):
                text_content = doc.page_content
            else:
                text_content = str(doc)

            try:
                base64.b64decode(text_content)
                b64.append(text_content)
            except Exception:
                text.append(text_content)

        return {"images": b64, "texts": text}

    def _build_prompt(self, kwargs):
        docs_by_type = kwargs["context"]
        user_question = kwargs["question"]

        context_text = ""
        if len(docs_by_type["texts"]) > 0:
            for text_element in docs_by_type["texts"]:
                if hasattr(text_element, 'text'):
                    context_text += text_element.text + "\n\n"
                else:
                    context_text += str(text_element) + "\n\n"

        prompt_template = f""""You are an expert document assistant specialized in providing accurate, contextual answers from multi-modal content including text, tables, charts, and images.
        Context: {context_text}
        
        Question: {user_question}

Please provide a comprehensive answer based on the available context. If the context includes visual elements that are relevant, reference them in your response.Provide only the direct answer to the question. No meta-commentary, explanations of your process, or statements about the context quality unless specifically relevant to answering the question."""
        prompt_content = [{"type": "text", "text": prompt_template}]

        # Add images to prompt if available
        if len(docs_by_type["images"]) > 0:
            for image in docs_by_type["images"]:
                prompt_content.append({
                    "type": "image_url",
                    "image_url": {"url": f"data:image/jpeg;base64,{image}"},
                })

        return ChatPromptTemplate.from_messages([HumanMessage(content=prompt_content)])

    def setup_rag_chain(self):
        vision_llm = ChatGroq(
            model="meta-llama/llama-4-scout-17b-16e-instruct",
            temperature=0.3,max_completion_tokens= 512
        )

        self.rag_chain = (
                {
                    "context": self.retriever | RunnableLambda(self._parse_docs),
                    "question": RunnablePassthrough(),
                }
                | RunnableLambda(self._build_prompt)
                | vision_llm
                | StrOutputParser()
        )

        self.rag_chain_with_sources = {
                                          "context": self.retriever | RunnableLambda(self._parse_docs),
                                          "question": RunnablePassthrough(),
                                      } | RunnablePassthrough().assign(
            response=(
                    RunnableLambda(self._build_prompt)
                    | vision_llm
                    | StrOutputParser()
            )
        )

    def query(self, question: str, return_sources: bool = False):
        if return_sources:
            return self.rag_chain_with_sources.invoke(question)
        else:
            return self.rag_chain.invoke(question)

    def process_document(self, file_path: str):
        file_extension = os.path.splitext(file_path)[1].lower()

        if file_extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {file_extension}")

        print(f"Processing {file_extension} file: {file_path}")

        chunks = self.extract_document_elements(file_path, file_extension)

        texts, tables, images = self.separate_elements(chunks)

        text_summaries, table_summaries, image_summaries = self.generate_summaries(texts, tables, images)

        self.load_to_vectorstore(texts, tables, images, text_summaries, table_summaries, image_summaries)

        self.setup_rag_chain()

        print("Document processing complete! Ready for queries.")
        return {
            "texts": len(texts),
            "tables": len(tables),
            "images": len(images)
        }

    def display_images_from_base64(self, images_b64, title="Images used in context"):
        if not images_b64:
            print("No images to display")
            return

        num_images = len(images_b64)
        if num_images == 1:
            fig, ax = plt.subplots(1, 1, figsize=(10, 8))
            axes = [ax]
        else:
            cols = min(3, num_images)
            rows = (num_images + cols - 1) // cols
            fig, axes = plt.subplots(rows, cols, figsize=(15, 5 * rows))
            axes = axes.flatten() if num_images > 1 else [axes]

        for idx, img_b64 in enumerate(images_b64):
            try:
                img_data = base64.b64decode(img_b64)
                img = Image.open(io.BytesIO(img_data))
                axes[idx].imshow(img)
                axes[idx].axis('off')
                axes[idx].set_title(f'Image {idx + 1}')
            except Exception as e:
                print(f"Error displaying image {idx + 1}: {e}")
                axes[idx].text(0.5, 0.5, f'Error loading image {idx + 1}',
                               ha='center', va='center', transform=axes[idx].transAxes)
                axes[idx].axis('off')

        for idx in range(num_images, len(axes)):
            axes[idx].axis('off')

        plt.suptitle(title, fontsize=16)
        plt.tight_layout()
        plt.show()

    def clear_vectorstore(self):
        """Clear the vectorstore for processing new documents."""
        self.vectorstore = Chroma(
            collection_name="enhanced_multi_modal_rag",
            embedding_function=self.embeddings
        )
        self.store = InMemoryStore()
        self.retriever = MultiVectorRetriever(
            vectorstore=self.vectorstore,
            docstore=self.store,
            id_key=self.id_key,
        )